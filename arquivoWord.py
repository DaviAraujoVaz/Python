from docx import Document

doc = Document()

DISCIPLINA = str(input("Nome da Disciplina: "))
p = doc.add_paragraph("Disciplina: ")
run = p.add_run()
run.text = DISCIPLINA
run.underline = True
run.bold = True

CARGA_HORARIA = str(input("Carga Horária: "))
p = doc.add_paragraph("Com carga horária de ")
run = p.add_run()
run.text = CARGA_HORARIA
run.underline = True
run.bold = True
run = p.add_run()
run.text = " horas."

p = doc.add_paragraph("Conteúdo da disciplina:")

tab = doc.add_table(rows=1, cols=2)
tab.style = "Colorful Grid Accent 6"
cels = tab.rows[0].cells
cels[0].text = "Número"
cels[0].width = 5
cels[1].text = "Assunto"

for i in range(1, 5):
    ASSUNTO_DISCIPLINA = str(input(f"Assunto da disciplina-{i}: "))
    dados = tab.add_row().cells
    dados[0].text = str(i)
    dados[0].width = 5
    dados[1].text = ASSUNTO_DISCIPLINA

doc.save("Disciplina.docx")
